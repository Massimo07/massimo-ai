import os
import openai
from fpdf import FPDF
from docx import Document
from PIL import Image
import requests

# ---------- CONFIG INIZIALE ----------
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY") or "INSERISCI-LA-TUA-KEY-QUI"
openai.api_key = OPENAI_API_KEY
LOGO_PATH = "data/logo.png" if os.path.exists("data/logo.png") else None

CORSI = [
    "network_marketing", "office", "adobe", "ai_assistant", "ai_copywriting", "copywriting", "ai_video",
    "video_editing", "automation", "autostima", "autotraining", "avatar_ai", "canva", "funnel", "gamification",
    "google_suite", "leadership", "mentalita_milionario", "notion", "personal_brand", "plugin_automazioni",
    "podcast", "podcast_radio", "presentazione_aziendale", "presentazione_marketing", "presentazione_prodotti",
    "recruiting", "social_adv", "social", "tiktok", "time_management", "vendita", "vr_training",
    "web_master", "web_design", "webinar", "youtube",
    "lingua_inglese", "lingua_francese", "lingua_spagnola", "lingua_tedesca", "lingua_portoghese",
    "lingua_rumena", "lingua_russa", "lingua_araba", "lingua_cinese", "lingua_olandese", "lingua_greca",
    "lingua_turca", "lingua_polacca", "lingua_svedese", "lingua_norvegese", "lingua_danese", "lingua_finlandese",
    "lingua_serba", "lingua_slovena", "lingua_estone", "lingua_lituana", "lingua_lettonia", "lingua_ucraina", "lingua_bulgara", "lingua_croata", "lingua_ceca", "lingua_slovacca", "lingua_ungherese"
]
LINGUE = [
    "it","en","fr","de","es","pt","ro","bg","hr","hu","cs","sk","pl","el","nl","ru","tr",
    "da","fi","no","sv","uk","lt","lv","et","sl","sr","ar","zh"
]
LIVELLI = [1,2,3,4,5,6,7]


def crea_cartella(path):
    if not os.path.exists(path):
        os.makedirs(path)

# ---------- FUNZIONI IMMAGINI ----------

def genera_immagine_ai(prompt, out_path):
    dalle_response = openai.images.generate(
        model="dall-e-3",
        prompt=prompt,
        n=1,
        size="1024x1024"
    )
    img_url = dalle_response.data[0].url
    img_data = requests.get(img_url).content
    with open(out_path, 'wb') as handler:
        handler.write(img_data)
    return out_path

# ---------- GPT CORSO COMPLETO ----------

def crea_contenuto_corso(corso, livello, lingua):
    prompt = f"""Crea un corso professionale, dettagliato e completo su '{corso}' – Livello {livello} ({lingua}).
• Copertina (descrivi l’immagine)
• Indice dettagliato (10 moduli)
• Spiegazione modulo per modulo (approfondita, MAI banale)
• Box motivazione, esercizi, quiz, checklist, glossario, casi studio
• Prompt immagini AI per ogni modulo (scrivi [Immagine: ...])
• Nessuna semplificazione eccessiva, solo qualità top.
Restituisci il tutto in formato markdown professionale."""
    system = "Sei il miglior formatore business d'Europa. Scrivi come un formatore di Harvard."
    response = openai.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": prompt}
        ],
        max_tokens=4096
    )
    return response.choices[0].message.content.strip()

# ---------- CREAZIONE FILE ----------

def crea_pdf(titolo, testo, immagini, out_path):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", "B", 20)
    pdf.cell(0, 10, titolo, ln=True, align="C")
    pdf.ln(10)
    if LOGO_PATH:
        pdf.image(LOGO_PATH, 10, 10, 40)
    pdf.set_font("Arial", "", 12)
    for line in testo.split('\n'):
        if line.startswith("[Immagine:") and immagini:
            img_path = immagini.pop(0)
            pdf.ln(5)
            pdf.image(img_path, w=120)
            pdf.ln(5)
        else:
            pdf.multi_cell(0, 8, line)
    pdf.output(out_path, "F")

def crea_docx(titolo, testo, immagini, out_path):
    doc = Document()
    doc.add_heading(titolo, 0)
    if LOGO_PATH:
        doc.add_picture(LOGO_PATH)
    for line in testo.split('\n'):
        if line.startswith("[Immagine:") and immagini:
            img_path = immagini.pop(0)
            doc.add_picture(img_path)
        else:
            doc.add_paragraph(line)
    doc.save(out_path)

# ---------- GENERAZIONE AUTOMATICA MAX ----------

def main():
    for corso in CORSI:
        for livello in LIVELLI:
            for lingua in LINGUE:
                titolo = f"{corso.replace('_',' ').title()} - Livello {livello} [{lingua}]"
                print(f"Generazione: {titolo} ...")
                folder = f"data/corsi/{corso}/L{livello}/{lingua}"
                crea_cartella(folder)
                testo = crea_contenuto_corso(corso, livello, lingua)
                img_prompts = []
                lines = testo.split('\n')
                for line in lines:
                    if "Immagine:" in line:
                        prompt_img = line.split("Immagine:")[1].strip(" []")
                        img_prompts.append(prompt_img)
                img_paths = []
                for i, prompt_img in enumerate(img_prompts):
                    img_path = f"{folder}/img_{i+1}.png"
                    genera_immagine_ai(prompt_img + " in stile business, alta qualità", img_path)
                    img_paths.append(img_path)
                # File markdown (sempre incluso)
                with open(f"{folder}/{corso}_L{livello}_{lingua}.md", "w", encoding="utf-8") as f:
                    f.write(testo)
                # PDF e Word (altissima qualità, pronto)
                crea_pdf(titolo, testo, img_paths.copy(), f"{folder}/{corso}_L{livello}_{lingua}.pdf")
                crea_docx(titolo, testo, img_paths.copy(), f"{folder}/{corso}_L{livello}_{lingua}.docx")
                print(f"✅ Creato {titolo} in {folder}")

if __name__ == "__main__":
    main()
