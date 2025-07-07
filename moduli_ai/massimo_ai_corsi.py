import os
from docx import Document
from docx.shared import Inches

# Lista delle lingue ufficiali + extra
lingue = [
    ("it", "Italiano"), ("en", "Inglese"), ("fr", "Francese"), ("de", "Tedesco"),
    ("es", "Spagnolo"), ("pt", "Portoghese"), ("nl", "Olandese"), ("el", "Greco"),
    ("ro", "Rumeno"), ("hu", "Ungherese"), ("pl", "Polacco"), ("cs", "Ceco"),
    ("sk", "Slovacco"), ("bg", "Bulgaro"), ("hr", "Croato"), ("sl", "Sloveno"),
    ("sv", "Svedese"), ("da", "Danese"), ("fi", "Finlandese"), ("et", "Estone"),
    ("lv", "Lettone"), ("lt", "Lituano"), ("no", "Norvegese"), ("ga", "Irlandese"),
    ("mt", "Maltese"), ("sq", "Albanese"), ("sr", "Serbo"), ("bs", "Bosniaco"),
    ("me", "Montenegrino"), ("mk", "Macedone"), ("tr", "Turco"), ("uk", "Ucraino"),
    ("be", "Bielorusso"), ("ka", "Georgiano"), ("hy", "Armeno"), ("rm", "Rumantsch"),
    ("lb", "Lussemburghese"), ("is", "Islandese"), ("ca", "Catalano"),
    ("ru", "Russo"), ("zh", "Cinese"), ("ar", "Arabo")
]

# Corsi tecnici
corsi_tecnici = [
    "adobe", "office", "ai_assistant", "ai_copywriting", "copywriting",
    "ai_video", "video_editing", "automation", "autostima", "autotraining",
    "avatar_ai", "canva", "funnel", "gamification", "google_suite",
    "leadership", "mentalita_milionario", "network_marketing", "notion",
    "personal_brand", "plugin_automazioni", "podcast", "podcast_radio",
    "presentazione_aziendale", "presentazione_piano_marketing",
    "presentazione_prodotti", "recruiting", "social_adv", "social",
    "tiktok", "time_management", "vendita", "vr_training", "web_master",
    "web_design", "webinar", "youtube"
]

livelli = [
    "livello1", "livello2", "livello3", "livello4",
    "livello5", "livello6", "livello7"
]

# Immagine placeholder (metti qui il percorso dell’immagine che vuoi usare, oppure lascia "")
img_placeholder = ""  # es: "immagini/placeholder.jpg"

base_tecnici = os.path.join("massimo_ai", "data", "corsi")
base_lingue = os.path.join("massimo_ai", "data", "corsi_lingue")

# Crea DOCX per tutti i corsi tecnici
for codice, nome_lingua in lingue:
    for corso in corsi_tecnici:
        for livello in livelli:
            folder = os.path.join(base_tecnici, codice, corso)
            os.makedirs(folder, exist_ok=True)
            filename = f"{livello}.docx"
            path = os.path.join(folder, filename)
            doc = Document()
            doc.add_heading(f"Corso di {corso.replace('_',' ').title()}", 0)
            doc.add_paragraph(f"Lingua: {nome_lingua}")
            doc.add_paragraph(f"Livello: {livello.title()}")
            doc.add_paragraph("Contenuto segnaposto: inserisci qui i materiali di alta professionalità.")
            if img_placeholder and os.path.exists(img_placeholder):
                doc.add_picture(img_placeholder, width=Inches(4))
            doc.save(path)

# Crea DOCX per tutti i corsi di lingua
for codice, nome_lingua in lingue:
    folder = os.path.join(base_lingue, codice)
    os.makedirs(folder, exist_ok=True)
    for livello in livelli:
        filename = f"{livello}.docx"
        path = os.path.join(folder, filename)
        doc = Document()
        doc.add_heading(f"Corso di {nome_lingua}", 0)
        doc.add_paragraph(f"Lingua: {nome_lingua}")
        doc.add_paragraph(f"Livello: {livello.title()}")
        doc.add_paragraph("Contenuto segnaposto: inserisci qui i materiali di alta professionalità per la lingua.")
        if img_placeholder and os.path.exists(img_placeholder):
            doc.add_picture(img_placeholder, width=Inches(4))
        doc.save(path)

print("✅ Tutti i file DOCX generati per corsi tecnici e di lingua in tutte le lingue e livelli!")
