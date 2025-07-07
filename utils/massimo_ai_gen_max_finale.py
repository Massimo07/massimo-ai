import os
import openai
from fpdf import FPDF
from docx import Document
import requests
from dotenv import load_dotenv

# ========== CONFIG BASE ==========
load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
openai.api_key = OPENAI_API_KEY
LOGO_PATH = "data/logo.png" if os.path.exists("data/logo.png") else None

CORSI = [
    "network_marketing", "office", "canva", "ai_copywriting", "leadership", "personal_brand", "vendita",
    "video_editing", "webinar", "ai_assistant", "copywriting", "ai_video", "automation", "autostima",
    "autotraining", "avatar_ai", "funnel", "gamification", "google_suite", "mentalita_milionario", "notion",
    "plugin_automazioni", "podcast", "podcast_radio", "presentazione_aziendale", "presentazione_marketing",
    "presentazione_prodotti", "recruiting", "social_adv", "social", "tiktok", "time_management", "vr_training",
    "web_master", "web_design", "youtube", "lingue"
]
# Tutte le lingue europee principali + russo, arabo, cinese, turco
LINGUE = [
    "it", "en", "fr", "de", "es", "pt", "nl", "el", "pl", "cs", "sk", "ro", "bg", "hr", "hu", "ru", "ar", "zh",
    "fi", "sv", "no", "da", "et", "lv", "lt", "sl", "sr", "uk", "mk", "ka", "sq", "tr", "ja"
]
LIVELLI = [1, 2, 3, 4, 5, 6, 7]

# ========== UTILS ==========
def crea_cartella(path):
    os.makedirs(path, exist_ok=True)

def genera_immagine_ai(prompt, out_path):
    try:
        dalle_response = openai.images.generate(
            model="dall-e-3",
            prompt=prompt,
            n=1,
            size="1024x1024"
        )
        img_url = dalle_response.data[0].url
        img_data = requests.get(img_url).content
        with open(out_path, 'wb') as handler:
            handler.write(img_data)
        return out_path
    except Exception as e:
        print(f"Errore immagine AI: {e}")
        return None

def crea_contenuto_corso(corso, livello, lingua):
    # Prompt ancora più dettagliato (professionalità massima)
    prompt = f"""Crea un corso professionale COMPLETO, dettagliato, avanzato e impaginato su '{corso}' – Livello {livello} [{lingua}].
- Copertina: descrivi immagine e titolo.
- Indice con almeno 10-12 moduli.
- Ogni modulo: spiegazione lunga, esempi, box motivazionali, casi studio, esercizi pratici, quiz, checklist, glossario termini.
- Per ogni modulo, scrivi il prompt per generare 1 immagine AI a tema (scrivi: [Immagine: ...]).
- Imposta tutto come un corso pronto da vendere a professionisti, MAI banale o corto, MAI riduttivo.
- Se corso = 'lingue', imposta il contenuto come corso di lingua utile per chi fa business, 7 livelli dalla base a madrelingua, full business.
Restituisci tutto in markdown professionale.
"""
    system = "Sei il migliore creatore di corsi business del mondo, pensi in modo avanzato e impagini come una casa editrice di Harvard."
    try:
        response = openai.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": prompt}
            ],
            max_tokens=4096
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"Errore GPT corso: {e}")
        return "Errore nella generazione del corso. Riprova."

def crea_pdf(titolo, testo, immagini, out_path):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", "B", 20)
    pdf.cell(0, 10, titolo, ln=True, align="C")
    pdf.ln(10)
    if LOGO_PATH:
        try:
            pdf.image(LOGO_PATH, 10, 10, 40)
        except:
            pass
    pdf.set_font("Arial", "", 12)
    for line in testo.split('\n'):
        if line.startswith("[Immagine:") and immagini:
            img_path = immagini.pop(0)
            if img_path and os.path.exists(img_path):
                pdf.ln(5)
                pdf.image(img_path, w=120)
                pdf.ln(5)
        else:
            pdf.multi_cell(0, 8, line)
    pdf.output(out_path, "F")

def crea_docx(titolo, testo, immagini, out_path):
    doc = Document()
    doc.add_heading(titolo, 0)
    if LOGO_PATH and os.path.exists(LOGO_PATH):
        try:
            doc.add_picture(LOGO_PATH)
        except:
            pass
    for line in testo.split('\n'):
        if line.startswith("[Immagine:") and immagini:
            img_path = immagini.pop(0)
            if img_path and os.path.exists(img_path):
                try:
                    doc.add_picture(img_path)
                except:
                    pass
        else:
            doc.add_paragraph(line)
    doc.save(out_path)

# ========== MAIN MAX ==========
def main():
    for corso in CORSI:
        for livello in LIVELLI:
            for lingua in LINGUE:
                titolo = f"{corso.replace('_',' ').title()} - Livello {livello} [{lingua}]"
                print(f"Generazione: {titolo} ...")
                folder = f"data/corsi/{corso}/L{livello}/{lingua}"
                crea_cartella(folder)
                testo = crea_contenuto_corso(corso, livello, lingua)
                img_prompts = []
                lines = testo.split('\n')
                for line in lines:
                    if "Immagine:" in line:
                        prompt_img = line.split("Immagine:")[1].strip(" []")
                        img_prompts.append(prompt_img)
                img_paths = []
                for i, prompt_img in enumerate(img_prompts):
                    img_path = f"{folder}/img_{i+1}.png"
                    genera_immagine_ai(prompt_img + " in stile business, alta qualità", img_path)
                    img_paths.append(img_path)
                # Salva markdown sempre
                with open(f"{folder}/{corso}_L{livello}_{lingua}.md", "w", encoding="utf-8") as f:
                    f.write(testo)
                # PDF e DOCX
                crea_pdf(titolo, testo, img_paths.copy(), f"{folder}/{corso}_L{livello}_{lingua}.pdf")
                crea_docx(titolo, testo, img_paths.copy(), f"{folder}/{corso}_L{livello}_{lingua}.docx")
                print(f"✅ Creato {titolo} in {folder}")

if __name__ == "__main__":
    main()
